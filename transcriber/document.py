"""
Document generation and formatting for Word (.docx) and plain text (.txt).
"""

from pathlib import Path
from typing import List, Tuple
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .audio import AudioSlice


def save_document(
    output_docx_path: Path,
    title: str,
    source_filename: str,
    total_duration_str: str,
    language: str,
    model_name: str,
    timestamp_interval: int,
    slices_data: List[Tuple[AudioSlice, str]],
    save_txt: bool = False
):
    """Generate a structured Word (.docx) document and optional text file."""
    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    
    # 1. Main Document Title
    h1 = doc.add_heading(title, level=1)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if h1.runs:
        h1.runs[0].font.size = Pt(20)
        h1.runs[0].font.bold = True
    
    # 2. Metadata Information Block
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = meta_p.add_run(f"Source File: {source_filename} | Total Duration: {total_duration_str}\n")
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(100, 100, 100)
    
    interval_desc = f"{timestamp_interval} mins" if timestamp_interval > 0 else "None"
    r2 = meta_p.add_run(f"Language: {language} | Model: {model_name} | Markers: {interval_desc} | Parts: {len(slices_data)}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # Spacer
    
    # 3. Add Sections for Each Slice
    for slice_info, transcript_text in slices_data:
        if slice_info.total_parts > 1:
            h2 = doc.add_heading(
                f"Part {slice_info.part_num:02d} ({slice_info.start_time_str} - {slice_info.end_time_str})",
                level=2
            )
            if h2.runs:
                h2.runs[0].font.size = Pt(14)
        
        paragraphs = transcript_text.strip().split("\n\n")
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            p = doc.add_paragraph(para)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(8)

    doc.save(str(output_docx_path))
    try:
        print(f"\n📄 Saved Word Document: {output_docx_path}")
    except Exception:
        print(f"\n[Saved Word Document] {output_docx_path}")

    # Optional plain text output
    if save_txt:
        txt_path = output_docx_path.with_suffix(".txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(f"{title}\n")
            f.write(f"Source: {source_filename} ({total_duration_str})\n")
            f.write("=" * 60 + "\n\n")
            for slice_info, transcript_text in slices_data:
                if slice_info.total_parts > 1:
                    f.write(f"\n--- Part {slice_info.part_num:02d} ({slice_info.start_time_str} - {slice_info.end_time_str}) ---\n\n")
                f.write(transcript_text.strip() + "\n\n")
        try:
            print(f"📝 Saved Plain Text:    {txt_path}")
        except Exception:
            print(f"[Saved Plain Text] {txt_path}")
